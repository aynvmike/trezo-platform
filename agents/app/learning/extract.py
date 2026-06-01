"""Trade extraction from non-CSV documents.

Supports four input shapes:
- PDF        -> sent to Claude with type=document (base64) for vision-grade
               extraction. Works for broker statements, trade confirmations,
               and most printable docs.
- image/*    -> sent to Claude as type=image (base64). Screenshots of trade
               logs, broker apps, P&L pages.
- xlsx       -> parsed locally with openpyxl when available, converted to CSV
               text, then sent to Claude for normalization.
- docx       -> parsed locally with python-docx when available, plain text
               sent to Claude.

The extractor asks Claude to return JSON in a specific shape: a list of
trade rows matching the import API's required + recommended columns.
The Python side wraps the call so the web layer just receives clean rows.

Anthropic SDK is imported lazily; missing key returns a 'not configured'
error so the route can surface a clean message instead of crashing.
"""

from __future__ import annotations

import base64
import io
import json
from typing import Any, Optional

import structlog

from app.config import get_settings

log = structlog.get_logger("trezo.learning.extract")


EXTRACTION_PROMPT = """\
You are extracting trade history from a document a user uploaded.

Return ONLY a JSON object with this exact shape:
{
  "rows": [
    {
      "ticker": "AAPL",
      "side": "long" | "short",
      "strategy": "<free text label; use 'manual' if unknown>",
      "entry_price": <number>,
      "exit_price": <number>,
      "quantity": <number>,
      "opened_at": "YYYY-MM-DD" or "YYYY-MM-DDTHH:MM:SSZ",
      "closed_at": "YYYY-MM-DD" or null,
      "realized_pnl_usd": <number> or null,
      "notes": "<short free text or null>"
    }
  ],
  "confidence": "high" | "medium" | "low",
  "notes": "<short explanation if anything is ambiguous>"
}

Rules:
- Only include CLOSED trades. Skip open positions.
- side: "long" = buy-then-sell. "short" = sell-then-buy.
- If you can't tell direction, default to "long".
- Skip rows you can't parse cleanly; do not invent prices.
- If realized_pnl_usd appears in the source, use it verbatim.
- If the document doesn't contain trade history at all, return {"rows": [], "confidence": "low", "notes": "no trades found"}.
- DO NOT include any prose, headers, or markdown. Only the JSON object.
"""


def _client():
    s = get_settings()
    if not s.anthropic_api_key:
        return None
    try:
        from anthropic import Anthropic
        return Anthropic(api_key=s.anthropic_api_key)
    except Exception:  # noqa: BLE001
        return None


def _parse_response(text: str) -> dict[str, Any]:
    """Pull JSON out of Claude's response, tolerating accidental
    markdown fences."""
    t = text.strip()
    if t.startswith("```"):
        t = t.strip("`")
        # drop language tag if present
        if "\n" in t:
            first, rest = t.split("\n", 1)
            if first.lower() in ("json", "javascript"):
                t = rest
    # Find first { and last } as a safety net
    a = t.find("{")
    b = t.rfind("}")
    if a >= 0 and b > a:
        t = t[a:b + 1]
    return json.loads(t)


async def extract_rows(
    *,
    content_type: str,
    file_bytes: bytes,
    filename: str = "",
) -> dict[str, Any]:
    """Single entry point. Routes by content_type."""
    client = _client()
    if client is None:
        return {
            "ok": False,
            "error": "ANTHROPIC_API_KEY not configured on the agents service.",
        }

    ct = content_type.lower().strip()

    # --- PDF: send as document ---
    if ct == "application/pdf" or filename.lower().endswith(".pdf"):
        return await _extract_via_document(client, file_bytes, ct or "application/pdf")

    # --- Image: send as image ---
    if ct.startswith("image/"):
        return await _extract_via_image(client, file_bytes, ct)

    # --- XLSX: convert to CSV text, then send as text ---
    if (ct in ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
               "application/vnd.ms-excel")
            or filename.lower().endswith((".xlsx", ".xls"))):
        text = _xlsx_to_text(file_bytes)
        if not text:
            return {"ok": False,
                    "error": "Couldn't parse the spreadsheet. Install openpyxl or save the sheet as CSV."}
        return await _extract_via_text(client, text)

    # --- DOCX: extract plain text, then send as text ---
    if (ct == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or filename.lower().endswith(".docx")):
        text = _docx_to_text(file_bytes)
        if not text:
            return {"ok": False,
                    "error": "Couldn't parse the Word document. Install python-docx or paste the text as CSV."}
        return await _extract_via_text(client, text)

    # --- Plain text / CSV / unknown: try as text ---
    try:
        text = file_bytes.decode("utf-8", errors="replace")
    except Exception:  # noqa: BLE001
        return {"ok": False,
                "error": f"Unsupported content type: {ct or 'unknown'}"}
    return await _extract_via_text(client, text)


# ---- Claude callers -------------------------------------------------------

def _call_claude(client, content_blocks: list[dict]) -> str:
    msg = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=4000,
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": EXTRACTION_PROMPT},
                *content_blocks,
            ],
        }],
    )
    # Sonnet returns a list of content blocks; pull the text.
    parts = []
    for c in msg.content:
        t = getattr(c, "text", None)
        if t:
            parts.append(t)
    return "\n".join(parts)


async def _extract_via_document(client, blob: bytes, media_type: str) -> dict[str, Any]:
    import asyncio
    b64 = base64.standard_b64encode(blob).decode()
    block = {
        "type": "document",
        "source": {
            "type": "base64",
            "media_type": "application/pdf",
            "data": b64,
        },
    }

    def _run():
        return _call_claude(client, [block])

    try:
        text = await asyncio.to_thread(_run)
        parsed = _parse_response(text)
        return {"ok": True, **parsed}
    except Exception as e:  # noqa: BLE001
        log.warning("extract.document_failed", error=str(e)[:300])
        return {"ok": False, "error": f"Claude PDF extraction failed: {str(e)[:200]}"}


async def _extract_via_image(client, blob: bytes, media_type: str) -> dict[str, Any]:
    import asyncio
    b64 = base64.standard_b64encode(blob).decode()
    block = {
        "type": "image",
        "source": {
            "type": "base64",
            "media_type": media_type,
            "data": b64,
        },
    }

    def _run():
        return _call_claude(client, [block])

    try:
        text = await asyncio.to_thread(_run)
        parsed = _parse_response(text)
        return {"ok": True, **parsed}
    except Exception as e:  # noqa: BLE001
        log.warning("extract.image_failed", error=str(e)[:300])
        return {"ok": False, "error": f"Claude image extraction failed: {str(e)[:200]}"}


async def _extract_via_text(client, text: str) -> dict[str, Any]:
    import asyncio
    if len(text) > 80_000:
        text = text[:80_000] + "\n\n[...truncated...]"

    block = {"type": "text", "text": text}

    def _run():
        return _call_claude(client, [block])

    try:
        out = await asyncio.to_thread(_run)
        parsed = _parse_response(out)
        return {"ok": True, **parsed}
    except Exception as e:  # noqa: BLE001
        log.warning("extract.text_failed", error=str(e)[:300])
        return {"ok": False, "error": f"Claude extraction failed: {str(e)[:200]}"}


# ---- Local converters -----------------------------------------------------

def _xlsx_to_text(blob: bytes) -> Optional[str]:
    """Convert XLSX bytes to a CSV-ish text dump. Returns None when the
    optional `openpyxl` library isn't installed."""
    try:
        import openpyxl  # type: ignore
    except ImportError:
        return None
    try:
        wb = openpyxl.load_workbook(io.BytesIO(blob), data_only=True)
        lines: list[str] = []
        for ws in wb.worksheets:
            lines.append(f"# Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [
                    (str(c).strip() if c is not None else "")
                    for c in row
                ]
                if any(cells):
                    lines.append(",".join(cells))
        return "\n".join(lines)
    except Exception as e:  # noqa: BLE001
        log.warning("extract.xlsx_failed", error=str(e)[:200])
        return None


def _docx_to_text(blob: bytes) -> Optional[str]:
    """Convert DOCX bytes to plain text. Returns None when the optional
    `python-docx` library isn't installed."""
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return None
    try:
        doc = Document(io.BytesIO(blob))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(",".join(cells))
        return "\n".join(parts)
    except Exception as e:  # noqa: BLE001
        log.warning("extract.docx_failed", error=str(e)[:200])
        return None
